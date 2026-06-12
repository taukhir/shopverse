from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/Shopverse_Centralized_Logging_And_Observability.docx")

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
INK = "202124"
WHITE = "FFFFFF"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E6F4EA"
PALE_YELLOW = "FFF4CE"
PALE_RED = "FCE8E6"


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Shopverse Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.2)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("SHOPVERSE  |  CENTRALIZED LOGGING AND OBSERVABILITY"), size=8.2, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Implementation and operations guide"), size=8.2, color=MUTED)


def title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("SHOPVERSE TECHNICAL GUIDE"), size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run("Centralized Logging and Observability"), size=27, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_font(
        p.add_run(
            "Spring Boot logging internals, SLF4J, Logback, Logstash JSON, MDC, "
            "Promtail, Loki, Micrometer, Prometheus, Grafana, Zipkin, queries, "
            "troubleshooting, retention, and production practices."
        ),
        size=12,
        color=MUTED,
    )
    table = doc.add_table(rows=1, cols=3)
    table_geometry(table, [3120, 3120, 3120])
    for cell, (label, value) in zip(
        table.rows[0].cells,
        (
            ("LOGS", "Logback -> Promtail -> Loki"),
            ("METRICS", "Micrometer -> Prometheus"),
            ("TRACES", "Micrometer Tracing -> Zipkin"),
        ),
    ):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(label), size=8, color=BLUE, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=9.5, color=NAVY, bold=True)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def callout(doc, heading, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(heading), size=10.2, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def code(doc, text, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(caption), size=8.8, color=MUTED, italic=True)
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.style = doc.styles["Shopverse Code"]
    for idx, line in enumerate(text.strip("\n").splitlines()):
        if idx:
            p.add_run().add_break()
        set_font(p.add_run(line), name="Consolas", size=8.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def grid(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table_geometry(table, widths)
    repeat_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), size=9, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_query(doc, title_text, query, explanation):
    doc.add_heading(title_text, level=2)
    code(doc, query)
    doc.add_paragraph(explanation)


def build():
    doc = Document()
    configure(doc)
    title(doc)

    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This document explains both the generic logging concepts and the concrete Shopverse "
        "implementation. It distinguishes logs, metrics, and traces; shows how Spring Boot creates "
        "structured events; explains how context travels across HTTP, Feign, and Kafka boundaries; "
        "and provides operational queries for Grafana, Loki, and Prometheus."
    )
    callout(
        doc,
        "Current implementation",
        "Shopverse uses JSON structured logs in Logstash format. The word Logstash describes the JSON "
        "schema selected by Spring Boot's encoder; Shopverse does not run a Logstash server. Promtail "
        "collects the JSON and pushes it to Loki.",
        PALE_GREEN,
    )

    doc.add_heading("2. Three Observability Signals", level=1)
    grid(
        doc,
        ["Signal", "Question answered", "Shopverse path"],
        [
            ("Logs", "What happened and why?", "SLF4J -> Logback -> Promtail -> Loki -> Grafana"),
            ("Metrics", "How often, how much, and is it healthy?", "Micrometer -> Actuator -> Prometheus -> Grafana"),
            ("Traces", "Where did a distributed request spend time?", "Micrometer Tracing -> Zipkin -> Grafana/Zipkin"),
        ],
        [1600, 3100, 4660],
    )
    code(
        doc,
        """
Client
  |
  v
API Gateway ---- logs ----> JSON files/stdout ----> Promtail ----> Loki
  |                  |
  |                  +---- metrics endpoint <---- Prometheus
  |
  +---- trace/span context ----------------------> Zipkin

Grafana queries Loki for logs and Prometheus for metrics.
Zipkin displays distributed traces.
""",
        "High-level Shopverse observability data flow",
    )

    page_break(doc)
    doc.add_heading("3. Generic Logging Architecture", level=1)
    grid(
        doc,
        ["Component", "Responsibility"],
        [
            ("SLF4J", "Facade/API used by application code."),
            ("Lombok @Slf4j", "Generates a class-level SLF4J logger named log."),
            ("Logback", "Spring Boot's default logging implementation and event dispatcher."),
            ("Logger", "Creates an event at TRACE, DEBUG, INFO, WARN, or ERROR."),
            ("Appender", "Writes an accepted event to console, file, socket, or another destination."),
            ("Encoder", "Converts a Java logging event into text or structured JSON."),
            ("MDC", "Adds scoped contextual fields such as correlationId."),
            ("Collector", "Tails sources, parses records, batches, and forwards them."),
            ("Loki", "Indexes bounded labels and stores compressed log chunks."),
            ("Grafana", "Queries and visualizes logs and metrics; it is not the primary datastore."),
        ],
        [2150, 7210],
    )

    doc.add_heading("Dependencies", level=2)
    code(
        doc,
        """
plugins {
    id 'org.springframework.boot'
    id 'io.freefair.lombok'
}

implementation 'org.springframework.boot:spring-boot-starter-web'
implementation 'org.springframework.boot:spring-boot-starter-actuator'
implementation 'org.springframework.boot:spring-boot-starter-zipkin'
runtimeOnly 'io.micrometer:micrometer-registry-prometheus'
""",
    )
    doc.add_paragraph(
        "The web starter brings Spring Boot logging, SLF4J, and Logback transitively. Lombok generates "
        "the logger field. Actuator exposes operational endpoints, the Prometheus registry renders "
        "meters in Prometheus format, and the Zipkin starter exports tracing spans."
    )
    callout(
        doc,
        "Dependency rule",
        "Application code should log through SLF4J and use one compatible provider. Multiple SLF4J "
        "providers can produce startup warnings or unpredictable binding behavior.",
        PALE_YELLOW,
    )

    doc.add_heading("Writing useful application logs", level=2)
    code(
        doc,
        """
@Slf4j
@Service
public class InventoryService {

    public void reserve(String orderNumber, Long productId, int quantity) {
        log.info(
                "Inventory reservation started orderNumber={} productId={} quantity={}",
                orderNumber, productId, quantity
        );
    }
}
""",
    )
    doc.add_paragraph(
        "Parameterized logging avoids unnecessary string concatenation and keeps fields recognizable. "
        "Log business boundaries, decisions, outcomes, retries, and failures rather than every method entry."
    )

    page_break(doc)
    doc.add_heading("4. Spring Boot Logging Internal Flow", level=1)
    code(
        doc,
        """
JVM starts main()
    |
    v
SpringApplication prepares Environment
    |
    v
LoggingApplicationListener initializes LoggingSystem early
    |
    v
Logback configuration is discovered
    |
    v
logback-spring.xml resolves <springProperty> values
    |
    v
Loggers, encoders, appenders, filters, and rolling policies are created
    |
    v
Application bean logs through SLF4J
    |
    v
Logback creates ILoggingEvent and snapshots MDC
    |
    +--> StructuredLogEncoder --> stdout
    |
    `--> StructuredLogEncoder --> active rolling file
""",
    )
    doc.add_paragraph(
        "Logging starts before the complete application context exists so startup failures can be reported. "
        "`logback-spring.xml` is preferred over `logback.xml` because Spring Boot can process extensions "
        "such as `<springProperty>` after the Environment is prepared."
    )

    doc.add_heading("What happens for one log statement", level=2)
    for item in [
        "Lombok's generated logger calls the SLF4J API.",
        "SLF4J delegates to the Logback provider.",
        "Logback checks the effective logger level.",
        "If enabled, Logback creates an immutable logging event containing timestamp, level, logger, thread, message arguments, throwable, key/value pairs, and a snapshot of MDC.",
        "Each attached appender receives the event.",
        "StructuredLogEncoder serializes it to one JSON object per line.",
        "ConsoleAppender writes stdout; RollingFileAppender writes and rotates the active file.",
    ]:
        numbered(doc, item)

    doc.add_heading("Log levels", level=2)
    grid(
        doc,
        ["Level", "Use"],
        [
            ("TRACE", "Very detailed framework or algorithm diagnostics."),
            ("DEBUG", "Developer diagnostics, normally disabled in production."),
            ("INFO", "Meaningful state transitions, requests, and successful outcomes."),
            ("WARN", "Recoverable failure, rejection, retry, or degraded operation."),
            ("ERROR", "Operation failed and requires investigation or recovery."),
        ],
        [1500, 7860],
    )

    page_break(doc)
    doc.add_heading("5. Shopverse logback-spring.xml", level=1)
    code(
        doc,
        """
<include resource="org/springframework/boot/logging/logback/defaults.xml"/>

<springProperty scope="context" name="APP_NAME"
                source="spring.application.name" defaultValue="SHOPVERSE"/>
<springProperty scope="context" name="LOG_FILE"
                source="logging.file.name"
                defaultValue="logs/${APP_NAME}.log"/>
<springProperty scope="context" name="HEALTH_LOG_FILE"
                source="shopverse.logging.health-file"
                defaultValue="logs/${APP_NAME}-health.log"/>
<springProperty scope="context" name="STRUCTURED_FORMAT"
                source="logging.structured.format.file"
                defaultValue="logstash"/>
""",
        "Spring properties made available to Logback",
    )
    doc.add_paragraph(
        "`defaults.xml` provides Spring Boot logging conversion rules and defaults. `<springProperty>` "
        "reads the central Spring Environment, including Config Server and environment-variable values, "
        "and exposes them as Logback variables."
    )
    code(
        doc,
        """
<appender name="CONSOLE" class="ch.qos.logback.core.ConsoleAppender">
    <encoder class="org.springframework.boot.logging.logback.StructuredLogEncoder">
        <format>${STRUCTURED_FORMAT}</format>
    </encoder>
</appender>

<appender name="APP_FILE"
          class="ch.qos.logback.core.rolling.RollingFileAppender">
    <file>${LOG_FILE}</file>
    <rollingPolicy
        class="ch.qos.logback.core.rolling.SizeAndTimeBasedRollingPolicy">
        <fileNamePattern>${LOG_FILE}.%d{yyyy-MM-dd}.%i.gz</fileNamePattern>
        <maxFileSize>10MB</maxFileSize>
        <maxHistory>7</maxHistory>
        <totalSizeCap>256MB</totalSizeCap>
    </rollingPolicy>
    <encoder class="org.springframework.boot.logging.logback.StructuredLogEncoder">
        <format>${STRUCTURED_FORMAT}</format>
    </encoder>
</appender>
""",
        "Console and rolling application file",
    )
    grid(
        doc,
        ["Configuration", "Meaning"],
        [
            ("fileNamePattern", "Date plus size index; old files are gzip-compressed."),
            ("maxFileSize=10MB", "Roll within the same day when the active file reaches 10 MB."),
            ("maxHistory=7", "Retain up to seven completed date periods."),
            ("totalSizeCap=256MB", "Bound total eligible archive storage."),
            ("root level=INFO", "Accept INFO and higher unless a specific logger overrides it."),
        ],
        [2600, 6760],
    )

    doc.add_heading("Health-log separation", level=2)
    code(
        doc,
        """
<logger name="io.shopverse.health" level="INFO" additivity="false">
    <appender-ref ref="HEALTH_FILE"/>
</logger>

<root level="INFO">
    <appender-ref ref="CONSOLE"/>
    <appender-ref ref="APP_FILE"/>
</root>
""",
    )
    doc.add_paragraph(
        "`additivity=false` prevents health records from also flowing to the root application appenders. "
        "Core business services retain health files for three days with a 64 MB cap, while application "
        "files retain seven days with a 256 MB cap."
    )

    page_break(doc)
    doc.add_heading("6. JSON and Logstash Structured Logging", level=1)
    doc.add_paragraph(
        "Shopverse currently configures `logging.structured.format.console` and "
        "`logging.structured.format.file` as `logstash`. Spring Boot's StructuredLogEncoder emits "
        "one JSON object per event."
    )
    code(
        doc,
        """
logging:
  file:
    name: ${LOG_FILE:logs/${spring.application.name}.log}
  structured:
    format:
      console: ${STRUCTURED_LOG_FORMAT:logstash}
      file: ${STRUCTURED_LOG_FORMAT:logstash}
    json:
      add:
        application: ${spring.application.name}
        environment: ${APP_ENVIRONMENT:local}
""",
        "Central configuration",
    )
    code(
        doc,
        """
{
  "@timestamp": "2026-06-12T10:30:00.123Z",
  "level": "INFO",
  "application": "ORDER-SERVICE",
  "environment": "local",
  "logger_name": "io.shopverse.order.service.OrderServiceImpl",
  "thread_name": "http-nio-8083-exec-4",
  "message": "Order created orderNumber=ORD-1003",
  "correlationId": "checkout-123",
  "traceId": "6a1e660de4db49fe47911954296ecce5",
  "spanId": "1ee04f11149f6bee"
}
""",
        "Representative structured event",
    )
    callout(
        doc,
        "Logstash format does not require Logstash",
        "Logstash format is a JSON field convention. Shopverse sends that JSON directly through Promtail "
        "to Loki. The deployed pipeline is not ELK and does not include Elasticsearch or Kibana.",
        PALE_GREEN,
    )
    doc.add_heading("Why JSON is preferable", level=2)
    for item in [
        "Fields remain independently searchable instead of requiring fragile regular expressions.",
        "Collectors can extract timestamp, application, level, traceId, spanId, and correlationId consistently.",
        "Grafana can display fields and create links without parsing human-formatted prefixes.",
        "Schema validation and redaction policies are easier to automate.",
        "Stack traces and messages remain part of one structured record.",
    ]:
        bullet(doc, item)

    page_break(doc)
    doc.add_heading("7. MDC: Thread-Scoped Logging Context", level=1)
    doc.add_paragraph(
        "MDC means Mapped Diagnostic Context. Logback's SLF4J MDC adapter maintains a map associated "
        "with the current execution thread, conceptually using thread-local storage. Values written "
        "by one servlet worker are not automatically visible to another worker."
    )
    code(
        doc,
        """
Thread: http-nio-8083-exec-4
MDC map:
  correlationId -> checkout-123
  traceId       -> 6a1e660de4db49fe47911954296ecce5
  spanId        -> 1ee04f11149f6bee

log.info("Order created");

Logback snapshots the current MDC map into the logging event.
StructuredLogEncoder writes those values as JSON fields.
""",
    )
    doc.add_heading("Safe scope and cleanup", level=2)
    code(
        doc,
        """
try (MDC.MDCCloseable ignored =
             MDC.putCloseable("correlationId", correlationId)) {
    filterChain.doFilter(request, response);
}
""",
    )
    doc.add_paragraph(
        "`putCloseable` adds the key and returns an AutoCloseable scope. Try-with-resources removes "
        "the key after normal completion, early return, or exception. Cleanup is essential because "
        "servlet and Kafka container threads are reused for unrelated work."
    )
    code(
        doc,
        """
Request A uses worker-4
  -> MDC correlationId=A
  -> logs
  -> close MDC scope
  -> worker-4 returns to pool clean

Request B later uses worker-4
  -> MDC correlationId=B
  -> cannot inherit A
""",
    )
    callout(
        doc,
        "MDC limitation",
        "MDC is thread-scoped, not automatically request-scoped across every asynchronous model. "
        "@Async, CompletableFuture, custom executors, Reactor, and Kafka consumer threads require "
        "explicit context propagation or restoration.",
        PALE_YELLOW,
    )

    doc.add_heading("Shopverse reusable Kafka context", level=2)
    code(
        doc,
        """
public static void run(String correlationId, Runnable action) {
    try (MDC.MDCCloseable ignored =
                 MDC.putCloseable(CorrelationConstants.MDC_KEY, correlationId)) {
        action.run();
    }
}

@KafkaListener(topics = "${shopverse.kafka.topics.order-created}")
public void onOrderCreated(String payload) {
    OrderCreatedEvent event = readEvent(payload);
    CorrelationContext.run(
            event.correlationId(),
            () -> handleOrderCreated(event)
    );
}
""",
    )
    doc.add_paragraph(
        "The event carries the business correlation ID. The listener restores it on the Kafka "
        "consumer thread before executing business logic, and the scope removes it afterward."
    )

    page_break(doc)
    doc.add_heading("8. Correlation Propagation Across Shopverse", level=1)
    code(
        doc,
        """
Client sends X-Correlation-Id (optional)
        |
        v
Gateway reuses or creates ID and adds response header
        |
        v
Servlet filter stores correlationId in MDC
        |
        v
Feign RequestInterceptor reads MDC and adds X-Correlation-Id
        |
        v
Downstream filter restores same value in downstream MDC
        |
        v
Kafka event stores correlationId in payload
        |
        v
Kafka listener uses CorrelationContext.run(...)
""",
    )
    code(
        doc,
        """
@Bean
RequestInterceptor correlationIdRequestInterceptor() {
    return template -> {
        String correlationId = MDC.get(CorrelationConstants.MDC_KEY);
        if (correlationId != null && !correlationId.isBlank()) {
            template.header(
                    CorrelationConstants.HEADER_NAME,
                    correlationId
            );
        }
    };
}
""",
        "Feign propagation",
    )
    grid(
        doc,
        ["Identifier", "Owner", "Purpose"],
        [
            ("correlationId", "Shopverse/application", "Connect one business journey across HTTP, retries, and Kafka events."),
            ("traceId", "Micrometer tracing", "Connect spans belonging to one distributed trace."),
            ("spanId", "Micrometer tracing", "Identify one operation within a trace."),
        ],
        [1900, 2300, 5160],
    )
    doc.add_paragraph(
        "A correlation ID may survive asynchronous delays and retries that create separate traces. "
        "Do not treat correlationId and traceId as interchangeable."
    )

    doc.add_heading("SLF4J fluent structured fields", level=2)
    code(
        doc,
        """
log.atInfo()
        .addKeyValue("correlationId", correlationId)
        .addKeyValue("method", method)
        .addKeyValue("path", path)
        .addKeyValue("status", status)
        .addKeyValue("durationMs", durationMs)
        .log("Gateway request completed");
""",
    )
    doc.add_paragraph(
        "`log.atInfo()` creates an INFO event builder. `addKeyValue` adds independent structured "
        "fields, and `.log(...)` emits the event. MDC is best for context shared by many logs; "
        "fluent key/value fields are best for values specific to one event."
    )

    page_break(doc)
    doc.add_heading("9. How Promtail Ships Shopverse Logs", level=1)
    code(
        doc,
        """
Rolling JSON file / Docker stdout
        |
        v
Promtail source discovery
        |
        v
positions.yml remembers last byte offset
        |
        v
JSON pipeline extracts fields
        |
        v
Only bounded values become Loki labels
        |
        v
Records are batched and pushed to
http://loki:3100/loki/api/v1/push
""",
    )
    grid(
        doc,
        ["Promtail job", "Source", "Purpose"],
        [
            ("shopverse-service-volume-files", "/service-logs/*/*.log", "Canonical JSON application logs in Docker volumes."),
            ("local-service-log-files", "/workspace/*/logs/*.log", "Logs from locally run services."),
            ("shopverse-health-log-files", "*-health.log", "Health/probe logs isolated from business events."),
            ("docker-containers", "Docker socket", "Infrastructure container stdout and startup diagnostics."),
        ],
        [2700, 2600, 4060],
    )
    code(
        doc,
        """
pipeline_stages:
  - json:
      expressions:
        timestamp: '"@timestamp"'
        level: level
        application: application
        traceId: traceId
        spanId: spanId
        correlationId: correlationId
        message: message
  - labels:
      level:
      application:
  - timestamp:
      source: timestamp
      format: RFC3339Nano
""",
        "Current JSON parsing pipeline",
    )
    doc.add_paragraph(
        "Extraction makes fields available during pipeline processing. Only `level` and `application` "
        "are promoted to labels because they have bounded cardinality. Unique correlation IDs and "
        "trace IDs remain JSON fields and are filtered after `| json`."
    )
    callout(
        doc,
        "Avoid duplicate ingestion",
        "Shopverse application containers are dropped from the Docker discovery job because their "
        "rolling files are the canonical Loki source. Infrastructure containers continue to use stdout.",
        PALE_GREEN,
    )
    callout(
        doc,
        "Promtail lifecycle",
        "Promtail is the collector implemented in this POC. For a new production platform, migrate "
        "to Grafana Alloy or another supported organization-standard collector while preserving "
        "positions, labels, parsing, retry, and security behavior.",
        PALE_YELLOW,
    )

    page_break(doc)
    doc.add_heading("10. Loki Storage and Retention", level=1)
    doc.add_paragraph(
        "Loki receives timestamped log lines plus labels. It indexes labels rather than indexing "
        "the complete log body, and stores compressed log chunks on filesystem storage in this POC."
    )
    code(
        doc,
        """
schema: TSDB v13
filesystem chunks: /loki/chunks
filesystem rules:  /loki/rules
replication factor: 1
retention period: 168h (7 days)
Docker volume: loki-data
""",
    )
    grid(
        doc,
        ["Storage", "Current retention", "Independent behavior"],
        [
            ("Application files", "7 days / 256 MB for business logs", "Managed by Logback rolling policy."),
            ("Health files", "3 days / 64 MB", "Managed separately by Logback."),
            ("Loki", "168 hours", "Managed by Loki compactor and Docker volume."),
        ],
        [2500, 3100, 3760],
    )
    doc.add_paragraph(
        "Deleting service log files does not delete already-ingested Loki records. Removing the "
        "`loki-data` volume deletes the local centralized store. Production deployments should use "
        "replicated components and durable object storage."
    )

    doc.add_heading("Label cardinality", level=2)
    grid(
        doc,
        ["Good labels", "Keep as parsed fields"],
        [
            ("application, level, job, environment, container", "correlationId, traceId, userId, orderNumber, raw path"),
        ],
        [4680, 4680],
    )
    doc.add_paragraph(
        "Every distinct label set creates a Loki stream. Unique identifiers used as labels can create "
        "millions of streams and damage ingestion and query performance."
    )

    page_break(doc)
    doc.add_heading("11. Viewing Logs in Grafana and Loki", level=1)
    doc.add_heading("Grafana steps", level=2)
    for step in [
        "Open http://localhost:3000 and sign in using the configured local credentials.",
        "Select Explore from the left navigation.",
        "Choose the Loki datasource.",
        "Select the exact time range in which the request was executed.",
        "Paste a LogQL query and select Run query.",
        "Expand a JSON log line to inspect application, level, correlationId, traceId, spanId, logger, and message.",
        "Copy traceId to Zipkin or use the configured derived trace link.",
    ]:
        numbered(doc, step)

    add_query(doc, "All application logs", '{log_type="application"}', "Shows business application logs while excluding the dedicated health stream.")
    add_query(doc, "All available streams", '{job=~".+"}', "Useful as a broad discovery query when labels are unknown. Keep the time range narrow.")
    add_query(doc, "One service", '{application="ORDER-SERVICE"}', "Change the application label to AUTH-SERVICE, USER-SERVICE, INVENTORY-SERVICE, PAYMENT-SERVICE, API-GATEWAY, CONFIG-SERVER, or DISCOVERY-SERVER.")
    add_query(
        doc,
        "One correlation ID across services",
        """
{job=~"shopverse-service-volume-files|shopverse-local-files|docker-containers"}
| json
| correlationId="checkout-123"
""",
        "This is the primary query for reconstructing one Shopverse business journey.",
    )
    add_query(
        doc,
        "One trace ID",
        """
{job=~"shopverse-service-volume-files|shopverse-local-files|docker-containers"}
| json
| traceId="6a1e660de4db49fe47911954296ecce5"
""",
        "Use traceId when investigating a single distributed trace and then open the same trace in Zipkin.",
    )

    page_break(doc)
    doc.add_heading("12. LogQL Query Reference", level=1)
    add_query(
        doc,
        "Errors and warnings for one service",
        """
{application="PAYMENT-SERVICE"}
| json
| level=~"WARN|ERROR"
""",
        "Use a bounded application selector before parsing to minimize scanned data.",
    )
    add_query(
        doc,
        "All errors grouped visually by application",
        """
{log_type="application"}
| json
| level="ERROR"
""",
        "In Grafana, display the application field or add it to the result line format.",
    )
    add_query(
        doc,
        "SAGA and outbox activity",
        """
{log_type="application"}
| json
| message=~"(?i).*(saga|outbox).*"
""",
        "Useful when an order remains in an intermediate state.",
    )
    add_query(doc, "One order number", '{log_type="application"} |= "ORD-1003"', "Text search is appropriate for a stable business identifier that is not a Loki label.")
    add_query(doc, "Dedicated health logs", '{log_type="health"}', "Use this only when debugging health probes or scrape availability.")
    add_query(
        doc,
        "JSON parsing failures",
        """
{log_type="application"}
| json
| __error__!=""
""",
        "Detects lines that do not match the expected structured JSON representation.",
    )
    add_query(
        doc,
        "Error count over five minutes",
        """
sum by (application) (
  count_over_time(
    {log_type="application"}
    | json
    | level="ERROR"
    [5m]
  )
)
""",
        "This derives a metric from logs. Prefer native application counters for primary alerting when available.",
    )

    doc.add_heading("Direct Docker diagnostics", level=2)
    code(
        doc,
        """
docker compose logs --tail=100 order-service
docker compose logs -f order-service
docker compose logs --since=10m order-service inventory-service payment-service
docker compose logs --tail=200 promtail
docker compose logs --tail=200 loki
""",
    )

    page_break(doc)
    doc.add_heading("13. Metrics: Micrometer to Prometheus", level=1)
    code(
        doc,
        """
Application code
    |
    v
MeterRegistry creates Counter / Timer / Gauge
    |
    v
PrometheusMeterRegistry stores current meter state
    |
    v
/actuator/prometheus renders text exposition
    |
    v
Prometheus scrapes every 15 seconds
    |
    v
Prometheus TSDB stores timestamped labeled samples
    |
    v
Grafana runs PromQL against Prometheus
""",
    )
    code(
        doc,
        """
meterRegistry.counter(
        "shopverse.service.requests.logged",
        "service", "ORDER-SERVICE",
        "method", request.getMethod(),
        "status", String.valueOf(status),
        "outcome", outcome(status)
).increment();
""",
        "Shopverse request counter",
    )
    doc.add_paragraph(
        "A Counter is a cumulative value that only increases until the process restarts. Prometheus "
        "normalizes the meter name to `shopverse_service_requests_logged_total`. Query `rate` or "
        "`increase` to understand activity within a time range."
    )
    grid(
        doc,
        ["Metric type", "Use"],
        [
            ("Counter", "Requests, failures, SAGA transitions, DLT events, published outbox events."),
            ("Gauge", "Current queue depth, active connections, current memory use."),
            ("Timer/Histogram", "Request duration count, sum, and percentile buckets."),
            ("DistributionSummary", "Distribution of non-duration values such as payload sizes."),
        ],
        [2200, 7160],
    )
    callout(
        doc,
        "Metrics are not logs",
        "Prometheus should not receive correlationId, traceId, orderNumber, username, or raw URL as "
        "metric labels. Those values create unbounded time-series cardinality. Use Loki for exact "
        "request investigations.",
        PALE_RED,
    )

    doc.add_heading("Timer recommendation", level=2)
    code(
        doc,
        """
Timer.builder("shopverse.gateway.request.duration")
        .tags("method", method, "status", String.valueOf(status))
        .register(meterRegistry)
        .record(durationMs, TimeUnit.MILLISECONDS);
""",
    )
    doc.add_paragraph(
        "A Timer provides count and duration distribution. Histograms support p50, p95, and p99 "
        "queries, which are more useful for latency analysis than only counting completed requests."
    )

    page_break(doc)
    doc.add_heading("14. Prometheus and Grafana Metric Queries", level=1)
    add_query(doc, "All Shopverse scrape targets", 'up{job="shopverse-services"}', "`1` means the metrics endpoint was scraped successfully; it does not prove checkout correctness.")
    add_query(doc, "Unavailable services", 'up{job="shopverse-services"} == 0', "Use Prometheus Targets to inspect the scrape error.")
    add_query(
        doc,
        "Request rate by service",
        """
sum by (application) (
  rate(http_server_requests_seconds_count[5m])
)
""",
        "Requests per second averaged over five minutes.",
    )
    add_query(
        doc,
        "5xx rate by service",
        """
sum by (application) (
  rate(http_server_requests_seconds_count{status=~"5.."}[5m])
)
""",
        "Shows server-side failures without mixing expected client errors.",
    )
    add_query(
        doc,
        "p95 HTTP latency",
        """
histogram_quantile(
  0.95,
  sum by (le, application) (
    rate(http_server_requests_seconds_bucket[5m])
  )
)
""",
        "Requires histogram buckets, which Shopverse enables for HTTP server requests.",
    )
    add_query(
        doc,
        "Custom request counter",
        """
sum by (service, status, outcome) (
  rate(shopverse_service_requests_logged_total[5m])
)
""",
        "Shows rate from the explicit Shopverse logging filter counter.",
    )
    add_query(
        doc,
        "SAGA transitions",
        """
sum by (stage) (
  increase(shopverse_saga_transitions_total[15m])
)
""",
        "Use to identify failed or stalled business transitions.",
    )
    add_query(
        doc,
        "Outbox publication outcomes",
        """
sum by (outcome) (
  increase(shopverse_outbox_publish_total[15m])
)
""",
        "A rising failed outcome requires inspection of Kafka connectivity and pending outbox state.",
    )
    add_query(
        doc,
        "DLT events and replays",
        """
sum by (service) (
  increase(shopverse_kafka_dlt_events_total[1h])
)

sum by (service) (
  increase(shopverse_kafka_dlt_replays_total[1h])
)
""",
        "Compare unresolved poison messages with operator recovery activity.",
    )
    add_query(
        doc,
        "Inventory conflicts and expiry",
        """
sum by (reason) (
  increase(shopverse_inventory_reservation_conflicts_total[15m])
)

sum(
  increase(shopverse_inventory_reservations_expired_total[15m])
)
""",
        "Useful for the last-item race and abandoned reservation demonstrations.",
    )

    page_break(doc)
    doc.add_heading("15. Grafana Responsibilities and Workflow", level=1)
    doc.add_paragraph(
        "Grafana is the common investigation and visualization interface. It does not scrape Spring "
        "Boot services and does not replace Loki, Prometheus, or Zipkin. The configured datasources are:"
    )
    grid(
        doc,
        ["Datasource", "Internal URL", "Use"],
        [
            ("Prometheus", "http://prometheus:9090", "Metrics, SLOs, recording rules, and alerts."),
            ("Loki", "http://loki:3100", "Structured logs and LogQL."),
            ("Zipkin", "http://zipkin:9411", "Distributed traces."),
        ],
        [2000, 2900, 4460],
    )
    doc.add_heading("Provisioned dashboards", level=2)
    for item in [
        "Shopverse Observability Overview: targets, request rates, latency, and recent logs.",
        "Shopverse Commerce Operations: SAGA transitions, payments, reservation conflicts, expiry, outbox, and DLT signals.",
    ]:
        bullet(doc, item)
    doc.add_heading("Recommended incident workflow", level=2)
    for step in [
        "Use a Prometheus panel or alert to identify the affected time window and service.",
        "Open Loki Explore and filter by application, severity, and correlationId.",
        "Find the first warning or error rather than only the final exception.",
        "Copy traceId and inspect the Zipkin span timeline.",
        "Inspect durable business state: order timeline, outbox row, reservation, payment, and DLT record.",
        "Validate recovery and confirm metrics return to normal.",
    ]:
        numbered(doc, step)
    code(
        doc,
        """
Prometheus anomaly
        |
        v
Loki logs by service/correlationId
        |
        v
Zipkin trace by traceId
        |
        v
Order timeline + Outbox + DLT + database state
        |
        v
Root cause and controlled recovery
""",
    )

    page_break(doc)
    doc.add_heading("16. Starting and Checking the Stack", level=1)
    code(
        doc,
        """
docker compose up -d
docker compose ps

# Direct interfaces
Grafana:    http://localhost:3000
Prometheus: http://localhost:9090
Loki ready: http://localhost:3100/ready
Zipkin:     http://localhost:9411

# Verify a service's metrics directly
http://localhost:8083/actuator/prometheus
""",
    )
    doc.add_heading("Prometheus target check", level=2)
    for step in [
        "Open http://localhost:9090/targets.",
        "Find job shopverse-services.",
        "Confirm each service target is UP.",
        "Open a failed target and inspect the scrape error.",
        "Confirm `/actuator/prometheus` is exposed and reachable from the Prometheus container.",
    ]:
        numbered(doc, step)

    doc.add_heading("End-to-end correlation demonstration", level=2)
    code(
        doc,
        """
1. Send an authenticated checkout with:
   X-Correlation-Id: demo-checkout-1003
   Idempotency-Key: demo-checkout-1003

2. In Grafana Loki Explore:
   {log_type="application"}
   | json
   | correlationId="demo-checkout-1003"

3. Confirm Order, Inventory, and Payment records appear.
4. Copy a traceId and open it in Zipkin.
5. Inspect SAGA and outbox metrics in Prometheus/Grafana.
""",
    )

    page_break(doc)
    doc.add_heading("17. Troubleshooting Guide", level=1)
    grid(
        doc,
        ["Symptom", "Checks"],
        [
            ("No service logs in Loki", "Check source file/stdout, Promtail logs, mounts, positions, JSON parse, Loki readiness, and time range."),
            ("Only health logs visible", "Query log_type=application and verify application jobs exclude *-health.log."),
            ("Correlation query empty", "Confirm header/event value, use | json, widen time range, and inspect raw event fields."),
            ("Duplicate logs", "Verify application containers are dropped from Docker discovery and positions data was not removed."),
            ("Metrics missing", "Check /actuator/prometheus, Prometheus Targets, exact normalized name, and trigger lazy meter creation."),
            ("p95 panel empty", "Confirm histogram buckets exist and percentile histograms are enabled."),
            ("Trace missing", "Check tracing enabled, sampling probability, Zipkin endpoint, and exporter connectivity."),
            ("Slow Loki query", "Narrow time range and labels before JSON parsing; avoid broad regex and high-cardinality labels."),
        ],
        [2600, 6760],
    )
    doc.add_heading("Promtail diagnostics", level=2)
    code(
        doc,
        """
docker compose logs --tail=200 promtail
docker compose logs --tail=200 loki
docker compose exec promtail cat /tmp/positions.yml
""",
    )
    doc.add_paragraph(
        "A missing positions file can cause rereading and duplicates. A stale position, path mismatch, "
        "permission problem, invalid timestamp, or malformed JSON can prevent expected ingestion."
    )

    page_break(doc)
    doc.add_heading("18. Logging and Observability Best Practices", level=1)
    for item in [
        "Use stable event names and field names across services.",
        "Log meaningful state transitions and external outcomes, not every method entry.",
        "Use parameterized or fluent logging instead of string concatenation.",
        "Pass exceptions as throwable arguments so stack traces are preserved.",
        "Never log passwords, bearer tokens, cookies, private keys, card data, or sensitive request bodies.",
        "Validate and length-limit caller-provided correlation IDs; reject control characters.",
        "Use try-with-resources or finally blocks to clean MDC on reused threads.",
        "Propagate context explicitly across Feign, Kafka, executors, and reactive boundaries.",
        "Keep unique identifiers as JSON fields, not Loki or Prometheus labels.",
        "Choose one canonical collection path for each workload to avoid duplicate events.",
        "Separate health/probe noise when it hides business activity.",
        "Use bounded local and central retention and test rotation behavior.",
        "Protect Actuator, Grafana, Loki, Prometheus, and Zipkin from untrusted networks.",
        "Use metrics for alerting, logs for evidence, traces for latency, and durable domain state for business truth.",
        "Source-control dashboards, datasources, alerts, and collector configuration.",
        "Monitor the observability stack itself for dropped lines, rejected samples, storage usage, and query latency.",
    ]:
        numbered(doc, item)

    doc.add_heading("Production hardening", level=2)
    for item in [
        "Replace local Loki filesystem storage with durable object storage and replication.",
        "Migrate Promtail to Grafana Alloy or another supported collector.",
        "Persist collector positions and configure backpressure and resource limits.",
        "Secure transport, authentication, authorization, tenant isolation, and encryption.",
        "Add Alertmanager notification routing, ownership, escalation, silence, and runbook links.",
        "Reduce trace sampling from 100 percent according to traffic and diagnostic requirements.",
        "Define SLOs for availability, latency, checkout correctness, and outbox freshness.",
        "Add Kafka consumer lag and oldest-pending-outbox-age metrics and alerts.",
    ]:
        bullet(doc, item)

    doc.add_heading("19. Quick Reference", level=1)
    grid(
        doc,
        ["Question", "Tool/query"],
        [
            ("Show all business logs", 'Loki: {log_type="application"}'),
            ("Show Order Service logs", 'Loki: {application="ORDER-SERVICE"}'),
            ("Follow one business journey", "Loki: | json | correlationId=\"...\""),
            ("Follow one distributed trace", "Loki traceId query, then Zipkin"),
            ("Check all services are scraped", 'Prometheus: up{job="shopverse-services"}'),
            ("Measure request rate", "Prometheus: rate(http_server_requests_seconds_count[5m])"),
            ("Measure p95 latency", "Prometheus: histogram_quantile over HTTP buckets"),
            ("Inspect SAGA failures", "SAGA transition metrics plus Loki correlation logs"),
            ("Inspect Kafka recovery", "DLT metrics, persisted DLT record, outbox logs, and replay audit"),
        ],
        [3300, 6060],
    )

    doc.add_heading("20. Related Repository References", level=1)
    for item in [
        "docs/observability/LOGGING-GENERIC.md",
        "docs/observability/STRUCTURED-LOGGING.md",
        "docs/observability/MDC-GENERIC.md",
        "docs/observability/MDC-CORRELATION-TRACING.md",
        "docs/observability/PROMTAIL.md",
        "docs/observability/LOKI.md",
        "docs/observability/PROMETHEUS.md",
        "docs/observability/GRAFANA.md",
        "docs/observability/OBSERVABILITY.md",
        "observability/promtail.yml",
        "observability/prometheus.yml",
        "observability/loki.yml",
        "<service>/src/main/resources/logback-spring.xml",
    ]:
        bullet(doc, item)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
