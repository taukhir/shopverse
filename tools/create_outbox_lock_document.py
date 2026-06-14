from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("documentation/docs/reliability/Outbox-Database-Lock-and-Kafka-Solution.docx")

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_RED = "FCE8E6"
PALE_GREEN = "E6F4EA"
WHITE = "FFFFFF"
BLACK = "1F1F1F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(BLACK)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("SHOPVERSE  |  RELIABILITY ENGINEERING")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("Outbox Database Lock and Kafka Solution")
    set_run_font(run, size=8.5, color=MUTED)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("TECHNICAL EXPLANATION")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Outbox Workers Holding Database Locks While Waiting for Kafka")
    set_run_font(run, size=25, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(
        "Problem statement, failure mechanics, implemented solution, transaction boundaries, "
        "data flow, and recovery behavior in Shopverse."
    )
    set_run_font(run, size=12.5, color=MUTED)

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    values = [
        ("Scope", "Order, Inventory, Payment"),
        ("Pattern", "Transactional Outbox"),
        ("Delivery", "At least once"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label.upper())
        set_run_font(r, size=8, color=BLUE, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=NAVY, bold=True)


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code_text, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        set_run_font(r, size=9, color=MUTED, italic=True)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.style = doc.styles["Code Block"]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code_text.strip("\n").splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_numbered(doc, title, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{title}: ")
    r.bold = True
    p.add_run(text)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_comparison_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2600, 3380, 3380])
    headers = ["Concern", "Previous implementation", "Current implementation"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("Lock during Kafka wait", "Yes", "No"),
        ("DB connection during Kafka wait", "Yes", "No"),
        ("Concurrent publisher protection", "Long pessimistic lock", "Short lock plus PROCESSING claim"),
        ("Crash recovery", "No explicit claim recovery", "claimed_at timeout returns stale claims"),
        ("Transaction duration", "Depends on Kafka response", "Only the database update duration"),
        ("Kafka outage effect", "Can exhaust DB pool and block APIs", "Publishing slows without retaining DB locks"),
    ]
    for idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col, (cell, text) in enumerate(zip(cells, values)):
            set_cell_shading(cell, WHITE if idx % 2 == 0 else LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col > 0 and text in {"Yes", "No"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            color = "9B1C1C" if text == "Yes" else ("137333" if text == "No" else BLACK)
            set_run_font(r, size=9.3, color=color, bold=text in {"Yes", "No"})


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(section)

    add_title_block(doc)
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The original outbox publisher opened a database transaction, acquired a pessimistic "
        "write lock, and then waited for Kafka acknowledgement inside that transaction. A slow "
        "or unavailable Kafka broker therefore caused the service to retain a MySQL row lock, "
        "connection, transaction, Hibernate context, and worker thread for the entire network wait."
    )
    add_callout(
        doc,
        "Implemented correction",
        "Shopverse now claims an event in a short transaction, commits and releases the database "
        "lock, sends the immutable event snapshot to Kafka outside any database transaction, and "
        "then opens another short transaction to mark the event PUBLISHED or return it to PENDING.",
        PALE_GREEN,
    )

    add_page_break(doc)
    doc.add_heading("1. Why Shopverse Uses an Outbox", level=1)
    doc.add_paragraph(
        "Checkout changes persistent business state and also starts asynchronous SAGA processing. "
        "For example, Order Service must save an order and publish OrderCreatedEvent. MySQL and "
        "Kafka are separate systems, so one ordinary database transaction cannot atomically commit both."
    )
    add_code(
        doc,
        """
Database transaction
|-- INSERT INTO orders
`-- INSERT INTO outbox_events (status = PENDING)

COMMIT

Scheduled outbox worker later publishes the event to Kafka.
""",
        "Transactional outbox write flow",
    )
    doc.add_heading("The failure window without an outbox", level=2)
    add_code(
        doc,
        """
Order saved successfully
        |
        v
Application crashes before Kafka send
        |
        v
Order exists, but Inventory Service never receives OrderCreatedEvent
""",
    )
    doc.add_paragraph(
        "The outbox closes this database-to-Kafka failure window by storing the business change "
        "and the outgoing event in the same MySQL transaction."
    )

    doc.add_heading("2. Previous Publishing Implementation", level=1)
    add_code(
        doc,
        """
@Transactional(propagation = Propagation.REQUIRES_NEW)
public void publish(Long eventId) {
    OutboxEvent event =
            repository.findByIdForUpdate(eventId).orElse(null);

    if (event == null || event.getStatus() != OutboxStatus.PENDING) {
        return;
    }

    kafkaTemplate.send(
            event.getTopic(),
            event.getMessageKey(),
            event.getPayload()
    ).get(10, TimeUnit.SECONDS);

    event.markPublished();
}
""",
        "Simplified previous worker",
    )
    doc.add_paragraph(
        "Because @Transactional covered the entire method, the transaction did not commit until "
        "Kafka returned and markPublished completed."
    )

    add_page_break(doc)
    doc.add_heading("3. What the Database Lock Does", level=1)
    doc.add_paragraph(
        "The repository method applies PESSIMISTIC_WRITE. MySQL effectively executes a SELECT FOR "
        "UPDATE statement and prevents another transaction from modifying or acquiring the same "
        "write lock until the first transaction completes."
    )
    add_code(
        doc,
        """
@Lock(LockModeType.PESSIMISTIC_WRITE)
@Query("select event from OutboxEvent event where event.id = :id")
Optional<OutboxEvent> findByIdForUpdate(@Param("id") Long id);

-- Approximate SQL
SELECT *
FROM outbox_events
WHERE id = ?
FOR UPDATE;
""",
    )
    doc.add_heading("Why a claim mechanism is required", level=2)
    doc.add_paragraph(
        "Multiple replicas or scheduler executions can discover the same PENDING event. Without "
        "coordination, both workers could publish it. The lock is therefore useful; the defect was "
        "holding it while performing slow external network I/O."
    )
    add_code(
        doc,
        """
Worker A finds event 100 (PENDING) ----> publishes to Kafka
Worker B finds event 100 (PENDING) ----> publishes to Kafka

Result without coordination: duplicate publication.
""",
    )

    doc.add_heading("4. The Lock-While-Waiting Problem", level=1)
    add_code(
        doc,
        """
BEGIN DATABASE TRANSACTION
        |
        v
SELECT ... FOR UPDATE       <-- row lock acquired
        |
        v
KafkaTemplate.send(...).get(10 seconds)
        |                   <-- lock and DB connection remain occupied
        v
status = PUBLISHED
        |
        v
COMMIT                      <-- lock finally released
""",
    )
    doc.add_paragraph("During the Kafka wait, the service retained:")
    for item in [
        "A MySQL transaction and pooled database connection.",
        "A row-level write lock on the outbox event.",
        "A Hibernate persistence context.",
        "The outbox worker thread waiting for an external system.",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    doc.add_heading("5. Operational Impact", level=1)
    doc.add_paragraph(
        "Kafka response time is outside MySQL's control. Broker outages, leader elections, network "
        "latency, metadata refreshes, or overloaded partitions can extend the wait. Database "
        "resources then become coupled to Kafka availability."
    )
    add_code(
        doc,
        """
Connection pool size: 10

10 outbox workers
    -> each locks one row
    -> each waits for Kafka
    -> all 10 database connections remain occupied

Checkout request
    -> asks for a database connection
    -> waits or times out
""",
        "Example: Kafka degradation indirectly blocks customer-facing database work",
    )
    add_callout(
        doc,
        "Core engineering problem",
        "A slow external message broker should not be allowed to retain scarce relational database "
        "connections and locks. The database work must be short and independent of the network wait.",
        PALE_RED,
    )

    doc.add_heading("6. Implemented Three-Phase Solution", level=1)
    add_numbered(
        doc,
        "Claim",
        "Lock the event briefly, change PENDING to PROCESSING, capture claimed_at, and commit.",
    )
    add_numbered(
        doc,
        "Publish",
        "Send an immutable OutboxMessage to Kafka after the claim transaction has completed.",
    )
    add_numbered(
        doc,
        "Finalize",
        "Open a new short transaction and change PROCESSING to PUBLISHED or back to PENDING.",
    )
    add_code(
        doc,
        """
Short DB transaction        No DB transaction         Short DB transaction
--------------------        -----------------         --------------------
PENDING                     Kafka send                PROCESSING
  -> PROCESSING       --->  wait for broker    --->     -> PUBLISHED
set claimed_at              acknowledgement             or PENDING
COMMIT / unlock
""",
    )

    add_page_break(doc)
    doc.add_heading("7. Phase One: Claim in a Short Transaction", level=1)
    add_code(
        doc,
        """
public OutboxMessage claim(Long eventId) {
    return transactionTemplate.execute(status -> {
        OutboxEvent event =
                repository.findByIdForUpdate(eventId).orElse(null);

        if (event == null
                || event.getStatus() != OutboxStatus.PENDING) {
            return null;
        }

        event.claim();
        return OutboxMessage.from(event);
    });
}
""",
    )
    doc.add_paragraph(
        "TransactionTemplate makes the boundary explicit. When execute returns, Spring has committed "
        "the transaction and returned its connection to the pool."
    )
    add_code(
        doc,
        """
public void claim() {
    status = OutboxStatus.PROCESSING;
    claimedAt = Instant.now();
    publishAttempts++;
}
""",
        "Entity state transition",
    )
    doc.add_paragraph(
        "PROCESSING acts as a durable claim. Other workers can inspect the row after commit but will "
        "not claim it because they require status PENDING."
    )
    doc.add_heading("Why use an immutable message snapshot?", level=2)
    add_code(
        doc,
        """
record OutboxMessage(
        Long id,
        String aggregateId,
        String eventType,
        String topic,
        String messageKey,
        String payload,
        String correlationId
) {}
""",
    )
    doc.add_paragraph(
        "Kafka receives ordinary immutable data instead of a managed JPA entity. Publishing therefore "
        "does not require an open Hibernate session, transaction, or database connection."
    )

    add_page_break(doc)
    doc.add_heading("8. Phase Two: Publish Without a Database Lock", level=1)
    add_code(
        doc,
        """
private void publishRecord(OutboxMessage message) {
    try {
        var result = kafkaTemplate.send(
                message.topic(),
                message.messageKey(),
                message.payload()
        ).get(sendTimeoutSeconds, TimeUnit.SECONDS);

        markPublished(message.id());
    } catch (Exception exception) {
        markFailed(message.id(), exception);
    }
}
""",
    )
    add_callout(
        doc,
        "Important boundary",
        "claim() has already returned before KafkaTemplate.send is called. There is no active "
        "outbox database transaction, no locked row, and no retained database connection while "
        "the worker waits for Kafka.",
        PALE_GREEN,
    )

    doc.add_heading("9. Phase Three: Finalize in a Short Transaction", level=1)
    add_code(
        doc,
        """
public void markPublished(Long eventId) {
    transactionTemplate.executeWithoutResult(status ->
            repository.findByIdForUpdate(eventId)
                    .filter(event ->
                            event.getStatus() == OutboxStatus.PROCESSING)
                    .ifPresent(OutboxEvent::markPublished)
    );
}
""",
    )
    add_code(
        doc,
        """
public void markPublished() {
    status = OutboxStatus.PUBLISHED;
    publishedAt = Instant.now();
    claimedAt = null;
    lastError = null;
}
""",
    )
    doc.add_paragraph(
        "The status check prevents a stale or unrelated operation from finalizing an event that is "
        "no longer in the expected PROCESSING state."
    )

    add_page_break(doc)
    doc.add_heading("10. Failure and Retry Flow", level=1)
    add_code(
        doc,
        """
public void markFailed(Throwable exception) {
    lastError = extractError(exception);
    claimedAt = null;
    status = OutboxStatus.PENDING;
}

PENDING -> PROCESSING -> Kafka failure -> PENDING
                                      -> next scheduled attempt
""",
    )
    doc.add_paragraph(
        "The current focused change returns failed sends to PENDING. Retry limits, exponential "
        "backoff, and a terminal failed state are separate reliability improvements and were not "
        "mixed into this fix."
    )

    doc.add_heading("11. Crash Recovery with claimed_at", level=1)
    doc.add_paragraph(
        "A service can crash after committing PROCESSING but before it finalizes the event. The "
        "claimed_at timestamp allows another scheduler run to identify and release abandoned claims."
    )
    add_code(
        doc,
        """
worker.releaseStaleClaims(
        Instant.now().minusMillis(claimTimeoutMs)
);

PROCESSING + claimed_at older than 30 seconds
        |
        v
releaseStaleClaim()
        |
        v
PENDING and available for another attempt
""",
    )
    add_code(
        doc,
        """
shopverse:
  outbox:
    publish-delay-ms: 1000
    send-timeout-seconds: 10
    claim-timeout-ms: 30000
""",
        "Centralized configuration",
    )
    doc.add_paragraph(
        "The claim timeout is intentionally longer than the Kafka send timeout, reducing the chance "
        "that a healthy in-flight publisher is incorrectly treated as abandoned."
    )

    add_page_break(doc)
    doc.add_heading("12. Complete Data Flow", level=1)
    add_code(
        doc,
        """
Authenticated checkout request
        |
        v
Order database transaction
|-- save order
`-- save OutboxEvent(PENDING)
        |
        v
Commit both records atomically
        |
        v
Scheduled outbox publisher selects PENDING events
        |
        v
Short claim transaction
|-- SELECT FOR UPDATE
|-- PENDING -> PROCESSING
|-- set claimed_at and increment attempts
`-- COMMIT and release lock
        |
        v
KafkaTemplate.send(...) outside DB transaction
        |
        +-----------------------------+
        |                             |
        v                             v
Kafka acknowledgement             Kafka failure
        |                             |
        v                             v
Short finalize transaction        Short finalize transaction
PROCESSING -> PUBLISHED           PROCESSING -> PENDING
""",
    )

    doc.add_heading("13. Before and After", level=1)
    add_comparison_table(doc)

    doc.add_heading("14. Remaining Delivery Guarantee", level=1)
    doc.add_paragraph(
        "The implementation provides at-least-once publication. If Kafka accepts an event and the "
        "service crashes before marking it PUBLISHED, stale-claim recovery can publish it again. "
        "This is expected for this outbox design; Kafka consumers must process events idempotently."
    )
    add_callout(
        doc,
        "Final result",
        "Shopverse still uses a database lock to claim work safely, but the lock now protects only "
        "short database state transitions. Kafka latency no longer determines how long an outbox "
        "row or pooled database connection remains locked.",
        PALE_GREEN,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
